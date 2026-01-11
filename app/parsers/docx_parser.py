"""DOCX document parser using python-docx."""

from typing import BinaryIO
from docx import Document
from loguru import logger

from .base import BaseParser


class DOCXParser(BaseParser):
    """Parser for Microsoft Word documents."""

    def supports(self, filename: str) -> bool:
        """Check if file is a DOCX."""
        return filename.lower().endswith(".docx")

    def parse(self, file: BinaryIO) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file: Binary file object containing DOCX data

        Returns:
            Extracted text content
        """
        text_parts = []

        try:
            doc = Document(file)

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")

        if not text_parts:
            raise ValueError("No text content found in DOCX")

        return "\n".join(text_parts)
